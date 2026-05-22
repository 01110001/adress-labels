from argparse import ArgumentParser
from pathlib import Path

import openpyxl
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


REQUIRED_COLUMNS = ["street", "city", "state", "zip_code", "company_name"]


def set_cell_border(cell, color="BDBDBD", size="8"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)

    for edge in ("top", "left", "bottom", "right"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=120, start=150, bottom=120, end=150):
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)

    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_row_height(row, cm):
    tr_pr = row._tr.get_or_add_trPr()
    height = tr_pr.find(qn("w:trHeight"))
    if height is None:
        height = OxmlElement("w:trHeight")
        tr_pr.append(height)
    height.set(qn("w:val"), str(int(cm * 567)))
    height.set(qn("w:hRule"), "exact")


def add_run(paragraph, text, size=9, bold=False, color=None):
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    return run


def clear_cell(cell):
    for paragraph in list(cell.paragraphs):
        paragraph._element.getparent().remove(paragraph._element)


def text_value(value):
    return "" if value is None else str(value).strip()


def prompt_sender_lines():
    print("Entrez les informations de l'expéditeur. Laissez vide pour ignorer un champ.")
    prompts = [
        "Nom de l'expéditeur",
        "Adresse",
        "Ville",
        "Province ou état",
        "Code postal",
        "Pays",
    ]
    return [value for value in (input(f"{prompt}: ").strip() for prompt in prompts) if value]


def add_sender_block(cell, sender_lines):
    if not sender_lines:
        return

    paragraph = cell.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(0)
    add_run(paragraph, "EXPÉDITEUR", size=6.5, bold=True, color=(90, 90, 90))

    for line in sender_lines:
        paragraph = cell.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        add_run(paragraph, line, size=7.2, color=(75, 75, 75))


def add_address_block(cell, row, sender_lines=None):
    clear_cell(cell)
    set_cell_border(cell)
    set_cell_margins(cell)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

    add_sender_block(cell, sender_lines)

    paragraph = cell.add_paragraph()
    if sender_lines:
        paragraph.paragraph_format.space_before = Pt(5)
        paragraph.paragraph_format.space_after = Pt(1)
    else:
        paragraph.paragraph_format.space_after = Pt(2)
    add_run(paragraph, "DESTINATAIRE", size=7, bold=True, color=(45, 45, 45))

    city_line = f"{text_value(row['city'])}, {text_value(row['state'])} {text_value(row['zip_code'])}".strip()
    destination_lines = [
        (text_value(row["company_name"]), 11.5 if sender_lines else 12.5, True),
        (text_value(row["street"]), 10.5 if sender_lines else 11.5, False),
        (city_line, 10.5 if sender_lines else 11.5, False),
        ("Canada", 10.5 if sender_lines else 11.5, False),
    ]

    for text, size, bold in destination_lines:
        paragraph = cell.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        add_run(paragraph, text, size=size, bold=bold)


def read_rows(source_path, sheet_name):
    workbook = openpyxl.load_workbook(source_path, data_only=True)
    if sheet_name not in workbook.sheetnames:
        available = ", ".join(workbook.sheetnames)
        raise ValueError(f"Feuille introuvable: {sheet_name}. Feuilles disponibles: {available}")

    sheet = workbook[sheet_name]
    headers = [cell.value for cell in sheet[1]]
    missing = [column for column in REQUIRED_COLUMNS if column not in headers]
    if missing:
        raise ValueError(f"Colonnes manquantes dans l'Excel: {', '.join(missing)}")

    rows = []
    for values in sheet.iter_rows(min_row=2, values_only=True):
        if any(values):
            rows.append(dict(zip(headers, values)))
    return rows


def build_document(rows, sender_lines=None):
    document = Document()
    section = document.sections[0]
    section.page_width = Cm(21.59)
    section.page_height = Cm(27.94)
    section.top_margin = Cm(1.1)
    section.bottom_margin = Cm(1.1)
    section.left_margin = Cm(1.0)
    section.right_margin = Cm(1.0)

    style = document.styles["Normal"]
    style.font.name = "Arial"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    style.font.size = Pt(9)

    labels_per_page = 8
    for page_start in range(0, len(rows), labels_per_page):
        if page_start:
            document.add_section(WD_SECTION_START.NEW_PAGE)

        table = document.add_table(rows=4, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False

        for table_row in table.rows:
            set_row_height(table_row, 6.35)
            for cell in table_row.cells:
                cell.width = Cm(9.65)

        page_rows = rows[page_start : page_start + labels_per_page]
        for index, row in enumerate(page_rows):
            row_index, column_index = divmod(index, 2)
            add_address_block(table.cell(row_index, column_index), row, sender_lines)

        for index in range(len(page_rows), labels_per_page):
            row_index, column_index = divmod(index, 2)
            cell = table.cell(row_index, column_index)
            clear_cell(cell)
            set_cell_border(cell, color="E5E5E5", size="4")

    return document


def main():
    parser = ArgumentParser(description="Créer un document Word d'étiquettes d'adresses pour colis.")
    parser.add_argument("excel_path", type=Path, help="Chemin du fichier Excel source.")
    parser.add_argument("output_path", type=Path, help="Chemin du document Word à créer.")
    parser.add_argument("--sheet", default="Table1", help="Nom de la feuille Excel à lire. Défaut: Table1.")
    parser.add_argument(
        "--set-expediteur",
        action="store_true",
        help="Demander les informations de l'expéditeur et les ajouter sur chaque étiquette.",
    )
    args = parser.parse_args()

    sender_lines = prompt_sender_lines() if args.set_expediteur else None
    rows = read_rows(args.excel_path, args.sheet)
    document = build_document(rows, sender_lines)
    args.output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(args.output_path)
    print(f"Document créé: {args.output_path}")
    print(f"Adresses incluses: {len(rows)}")


if __name__ == "__main__":
    main()
